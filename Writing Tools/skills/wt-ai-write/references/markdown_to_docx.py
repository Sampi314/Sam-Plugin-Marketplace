"""
Convert a daily AI blog markdown file to a publication-ready .docx in the
exact Liam editorial format.

The output mirrors `From Liam/AI Blog 054 - Claude Skills Part 1.docx`:

- All paragraphs use the **Normal** style (Calibri 11 pt). Liam does NOT use
  Word's built-in Title / Heading 1 / Heading 2 styles; he formats headings
  with direct bold + (for the title) larger font size on Normal paragraphs.
- The title: Normal paragraph, single bold run at 18 pt.
- Section headings: Normal paragraph, bold run at default size.
- `TBC` placeholder: plain Normal paragraph.
- Body paragraphs: Normal, default formatting.
- Numbered and bulleted list items: **List Paragraph** style (not List Number
  or List Bullet). The lead phrase in a list item, if marked `**Lead:**` in
  the markdown, renders as bold + colon, the rest as normal weight —
  matching Liam's `Anthropic's official repository: start with...` cadence.
- Code blocks: monospace (Consolas 9.5 pt) with light grey shading.
- Block quotes: italic, indented (used for the "fast version" Claude prompt).
- Images: embedded inline, centred, sized to ~70 % page width.

The template `article-template.docx` (sibling file) provides Liam's exact
page size, margins, section properties, default fonts and style table.

Usage:
    python markdown_to_docx.py --input article.md --output article.docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    sys.stderr.write("Missing dependency: python-docx\n"
                     "Install with: pip install python-docx\n")
    sys.exit(1)


REF_DIR = Path(__file__).resolve().parent
DEFAULT_TEMPLATE = REF_DIR / "article-template.docx"

TITLE_SIZE = Pt(12)
HEADING_SIZE = Pt(11)
CODE_FONT = "Consolas"
CODE_SIZE = Pt(9.5)
INLINE_CODE_SIZE = Pt(10)
CODE_FILL = "F4F4F4"
PAGE_WIDTH_INCHES = 6.0  # approximate text width on A4 with Liam's margins

_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _shade(paragraph, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _inline(paragraph, text: str) -> None:
    """Parse **bold**, *italic*, `code` runs."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = CODE_FONT
            r.font.size = INLINE_CODE_SIZE
        else:
            paragraph.add_run(tok)


def _add_title(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = TITLE_SIZE


def _add_heading(doc, text: str) -> None:
    """Section heading — Normal paragraph, bold run, blank row before
    matching Liam's spacing convention (real empty paragraph as separator)."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = HEADING_SIZE


def _add_code_block(doc, code_text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(code_text)
    r.font.name = CODE_FONT
    r.font.size = CODE_SIZE
    _shade(p, CODE_FILL)


def _add_image(doc, base_dir: Path, src: str) -> None:
    img_path = (base_dir / src).resolve()
    if not img_path.exists():
        sys.stderr.write(f"  [warn] image not found: {img_path}\n")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Inches(PAGE_WIDTH_INCHES * 0.7))


def _add_list_item(doc, text: str) -> None:
    """List item in Liam's format: 'List Paragraph' style, lead **bold:** then normal."""
    p = doc.add_paragraph(style="List Paragraph")
    _inline(p, text)


def _add_blockquote(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    r = p.add_run(text)
    r.italic = True


def convert(md_path: Path, out_path: Path, template_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    base_dir = md_path.parent

    if template_path.exists():
        doc = Document(str(template_path))
        # Template contains one empty placeholder paragraph — remove it
        body = doc.element.body
        sectPr = body.find(qn("w:sectPr"))
        for child in list(body):
            tag = child.tag.split("}")[-1]
            if tag in ("p", "tbl"):
                body.remove(child)
        # We will re-add sectPr at the end if it was there
        if sectPr is not None:
            body.append(sectPr)
    else:
        sys.stderr.write(f"[warn] template not found at {template_path}, "
                         "falling back to blank document\n")
        doc = Document()

    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        # fenced code block
        if line.startswith("```"):
            if in_code:
                _add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # title
        if line.startswith("# "):
            _add_title(doc, stripped[2:].strip())
            i += 1
            continue

        # H2 (section heading)
        if line.startswith("## "):
            _add_heading(doc, stripped[3:].strip())
            i += 1
            continue

        # H3
        if line.startswith("### "):
            _add_heading(doc, stripped[4:].strip())
            i += 1
            continue

        # image
        img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img:
            _add_image(doc, base_dir, img.group(2))
            i += 1
            continue

        # numbered list (1. content)
        num = re.match(r"^(\d+)\.\s+(.*)", line)
        if num:
            _add_list_item(doc, num.group(2))
            i += 1
            continue

        # bullet list (- content) or (* content)
        bullet = re.match(r"^[-*]\s+(.*)", line)
        if bullet:
            _add_list_item(doc, bullet.group(1))
            i += 1
            continue

        # block quote (possibly multi-line)
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                content = lines[i][1:].lstrip()
                if content:
                    quote_lines.append(content)
                i += 1
            _add_blockquote(doc, " ".join(quote_lines))
            continue

        # TBC line
        if stripped == "TBC":
            doc.add_paragraph("TBC")
            i += 1
            continue

        # whole-paragraph italic: wrapped in single * (not **)
        # e.g. *Welcome back... directory.*  -> entire paragraph italic, inline
        # `**bold**` and `code` formatting inside still applies.
        if (stripped.startswith("*") and not stripped.startswith("**")
                and stripped.endswith("*") and not stripped.endswith("**")
                and len(stripped) > 2):
            p = doc.add_paragraph()
            _inline(p, stripped[1:-1])
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        _inline(p, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--template", default=DEFAULT_TEMPLATE, type=Path,
                        help=f"Path to article-template.docx (default: {DEFAULT_TEMPLATE}).")
    args = parser.parse_args()

    if not args.input.exists():
        sys.stderr.write(f"Input not found: {args.input}\n")
        return 1

    convert(args.input.resolve(), args.output.resolve(), args.template.resolve())
    print(f"Saved: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
