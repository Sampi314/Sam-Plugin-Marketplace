"""
General-purpose Markdown to Word (.docx) converter.

Usage:
    python scripts/md_to_docx.py scope.md                    # → scope.docx (same directory)
    python scripts/md_to_docx.py scope.md model_flow.md      # multiple files
    python scripts/md_to_docx.py scope.md -o output/          # specify output directory
    python scripts/md_to_docx.py *.md                         # all .md files in current dir

Handles: headings, tables, bold/italic, bullet lists, code blocks, horizontal rules.
Requires: python-docx (pip install python-docx)
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_md_to_blocks(md_text: str) -> list[dict]:
    """Parse markdown text into a list of typed blocks."""
    blocks = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            i += 1  # skip closing ```
            continue

        # Table (line with |)
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                text = re.sub(r"^\s*[-*]\s+", "", lines[i])
                items.append(text)
                i += 1
            blocks.append({"type": "list", "items": items})
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                text = re.sub(r"^\s*\d+\.\s+", "", lines[i])
                items.append(text)
                i += 1
            blocks.append({"type": "numbered_list", "items": items})
            continue

        # Blockquote
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            blocks.append({"type": "quote", "text": "\n".join(quote_lines)})
            continue

        # Regular paragraph (skip blank lines)
        if line.strip():
            para_lines = []
            while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6}\s|[-*]\s|\d+\.\s|```|---|\|)", lines[i]):
                para_lines.append(lines[i])
                i += 1
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
            continue

        i += 1

    return blocks


def add_formatted_text(paragraph, text: str):
    """Add text with bold/italic/code formatting to a paragraph."""
    # Split on formatting markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:
            paragraph.add_run(part)


def parse_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse markdown table lines into headers and rows."""
    def split_row(line):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    headers = split_row(table_lines[0])
    rows = []
    for line in table_lines[1:]:
        # Skip separator lines (---|---|---)
        if re.match(r"^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?\s*$", line):
            continue
        rows.append(split_row(line))
    return headers, rows


def convert_md_to_docx(md_path: Path, output_path: Path):
    """Convert a single .md file to .docx."""
    md_text = md_path.read_text(encoding="utf-8")
    blocks = parse_md_to_blocks(md_text)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = min(block["level"], 4)  # Word supports heading 1-4 reliably
            heading = doc.add_heading(level=level)
            add_formatted_text(heading, block["text"])

        elif btype == "paragraph":
            para = doc.add_paragraph()
            add_formatted_text(para, block["text"])

        elif btype == "list":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Bullet")
                add_formatted_text(para, item)

        elif btype == "numbered_list":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Number")
                add_formatted_text(para, item)

        elif btype == "quote":
            para = doc.add_paragraph()
            para.style = doc.styles["Quote"] if "Quote" in doc.styles else doc.styles["Normal"]
            para.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(para, block["text"])

        elif btype == "code":
            para = doc.add_paragraph()
            run = para.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            para.paragraph_format.left_indent = Inches(0.3)

        elif btype == "table":
            headers, rows = parse_table(block["lines"])
            num_cols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Headers
            for j, header in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(header)
                run.bold = True

            # Data rows
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j < num_cols:
                        cell = table.rows[i + 1].cells[j]
                        cell.text = ""
                        add_formatted_text(cell.paragraphs[0], cell_text)

        elif btype == "hr":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)

    doc.save(str(output_path))
    print(f"  {md_path.name} -> {output_path.name}")


def main():
    parser = argparse.ArgumentParser(description="Convert Markdown files to Word (.docx)")
    parser.add_argument("files", nargs="+", help="Markdown files to convert")
    parser.add_argument("-o", "--output-dir", help="Output directory (default: same as input file)")
    args = parser.parse_args()

    for file_str in args.files:
        md_path = Path(file_str)
        if not md_path.exists():
            print(f"  Skipping {file_str} — file not found")
            continue
        if md_path.suffix.lower() != ".md":
            print(f"  Skipping {file_str} — not a .md file")
            continue

        if args.output_dir:
            out_dir = Path(args.output_dir)
            out_dir.mkdir(parents=True, exist_ok=True)
            output_path = out_dir / md_path.with_suffix(".docx").name
        else:
            output_path = md_path.with_suffix(".docx")

        convert_md_to_docx(md_path, output_path)

    print("Done.")


if __name__ == "__main__":
    main()
